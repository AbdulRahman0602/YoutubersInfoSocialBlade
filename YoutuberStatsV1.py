import bs4
import requests
import re
import docx

def getinfo(name):
    #Connecting to website.
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
    }
    res = requests.get(name, headers=headers)
    res.raise_for_status()
    #Code of Extration of information.
    soup = bs4.BeautifulSoup(res.text, 'html.parser')
    #Basic info
    elems = soup.select('#socialblade-user-content > div:nth-child(1) > div:nth-child(2)')
    if elems:
        numbers = elems[0].text.strip()
        numRegex = re.compile(r'''
            # 53,000 or 199 or 1,947 or 55 or 2
                (
                   (\d{1,3})?
                   (,)?
                   (\d{3}|\d{2}|\d{1})           
                )           
        ''', re.VERBOSE)
        extractedNum = numRegex.findall(numbers)
        allNumbers = []

        for Numbers in extractedNum:
            allNumbers.append(Numbers[0])
            #result = '\n'.join(allNumbers)+'\n'
        #length = len(allNumbers)
        # for Num in range(0,length):
            # if Num == 0:
            #     print("Social Rank: "+allNumbers[Num])
                
            # elif Num == 1:
            #     print("Subscribe Rank: "+allNumbers[Num])
            # elif Num == 2:
            #     print("Video Views Rank: "+allNumbers[Num])
            # elif Num == 3:
            #     print("Country Rank: "+allNumbers[Num])
            # elif Num == 4:
            #     print("Domain Rank: "+allNumbers[Num])
    # else:
    #     print("Data not found")
        doc = docx.Document()
        doc.add_heading('Youtubers Data')
        text = (f""" 
                    Social Rank: {allNumbers[0]}
                    Subscribe Rank: {allNumbers[1]}
                    Country Rank: {allNumbers[2]}
                    Domain Rank: {allNumbers[3]}
                """)
        doc.add_paragraph(text)

    data = []
    #Getting the stats info
    for i in range(7,20):
        elems = soup.select(f'#socialblade-user-content > div:nth-child({i})')
        components = elems[0].text.split()
        # for i in range(0,9):
        #     print(components[i], end=" ")
        # print("\n")
        # print("********************************************************************************")
        entry={
            "Date" : components[0],
            "Day" : components[1],
            "Subscribers_change" : components[2],
            "Total_subscribers" : components[3],
            "Views_change" : components[4],
            "Total_views" : components[5],
            "Earningsl" : components[6],
            "Earningsu" : components[8]
        }
        data.append(entry)
        
        # print(f"""Date: {date},\nDay: {day},
        #       \nSubscribers Change: {subscribers_change}, \nTotal Subscribers: {total_subscribers},
        #        \nViews Change: {views_change},\nTotal Views: {total_views}, 
        #        \nEstimated Earnings: {estimated_earnings_low} - {estimated_earnings_high}""")
        # print("*****************************************************************")

        for entry in data:
            text = (f"""Date: {entry['Date']}
                        Day: {entry['Day']} 

                        Subscribers Change: {entry['Subscribers_change']} 
                        Total Subscribers: {entry['Total_subscribers']} 

                        Views Change: {entry['Views_change']}
                        Total Views: {entry['Total_views']} 

                        Estimated Earnings: {entry['Earningsl']} - {entry['Earningsu']}""")
        doc.add_paragraph(text)

    doc.save('D:\Computer Languages\Programming Languages\AutomateTheBooringStuffCodes\output.docx')
    
cname=input("")
getinfo(f'https://socialblade.com/youtube/c/{cname}_')



